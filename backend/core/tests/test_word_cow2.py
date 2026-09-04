"""Phase 1–2: Official COW 02 prints as filled Word; blanks from Word pages."""
from io import BytesIO

from django.contrib.auth.models import Group, User
from django.test import TestCase
from docx import Document
from docx.table import _Cell
from rest_framework.authtoken.models import Token
from rest_framework.test import APIClient

from core.form_atlas import ATLAS_FORMS, form_meta
from core.official_blanks import ATLAS_VERSION, blank_info
from core.word_forms import fill_cow2_docx, official_cow2_path

SAMPLE = {
    '__display.personnel': 'Kelebogile Maraba',
    'household.org_household_number': '472',
    'household.town': 'Westonaria',
}


class FillOfficialCow2Tests(TestCase):
    def test_template_present(self):
        self.assertTrue(official_cow2_path().exists())

    def test_identity_fills(self):
        doc = Document(BytesIO(fill_cow2_docx(SAMPLE)))
        header = doc.tables[0]
        self.assertEqual(_Cell(header.rows[0]._tr.tc_lst[2], header).text.strip(), '472')
        self.assertEqual(_Cell(header.rows[1]._tr.tc_lst[3], header).text.strip(), 'Westonaria')
        self.assertEqual(_Cell(doc.tables[2].rows[1]._tr.tc_lst[0], doc.tables[2]).text.strip(), 'Kelebogile Maraba')


class PrintCow2WordTests(TestCase):
    def setUp(self):
        Group.objects.get_or_create(name='admin')
        self.user = User.objects.create_user('cow2printer', password='x')
        self.user.groups.add(Group.objects.get(name='admin'))
        self.client = APIClient()
        self.client.credentials(HTTP_AUTHORIZATION=f'Token {Token.objects.create(user=self.user).key}')

    def test_print_cow2_downloads_docx_not_canvas(self):
        created = self.client.post('/api/households/', {'town': 'Westonaria'}, format='json')
        hid = created.data['id']
        token = Token.objects.get(user=self.user).key
        response = self.client.get(f'/api/print/cow2_note/?household_id={hid}&token={token}')
        self.assertEqual(response.status_code, 200)
        self.assertIn(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            response['Content-Type'],
        )
        self.assertNotIn(b'official-payload', response.content)


class Cow2WordAtlasTests(TestCase):
    def test_cow2_uses_word_geometry(self):
        self.assertEqual(ATLAS_FORMS['cow2_note']['geometry'], 'word')
        self.assertEqual(form_meta('cow2_note')['pages'], 2)
        info = blank_info('cow2_note', 0)
        self.assertEqual(info.get('source'), 'docx')
        self.assertIn('COW_2_Process_note_04042019.docx', info.get('source_docx') or '')
        self.assertIn('cow2', ATLAS_VERSION)
