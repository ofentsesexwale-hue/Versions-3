"""Remaining Official Word pack: C06, educational, process notes, CW 4a/4b/11/13."""
from io import BytesIO

from django.contrib.auth.models import Group, User
from django.test import TestCase
from docx import Document
from docx.table import _Cell
from rest_framework.authtoken.models import Token
from rest_framework.test import APIClient

from core.form_atlas import ATLAS_FORMS, form_meta, has_geometry
from core.official_blanks import ATLAS_VERSION, blank_info
from core.word_forms import (
    fill_c06_docx,
    fill_checklist_docx,
    fill_content_page_docx,
    fill_educational_docx,
    fill_exit_docx,
    fill_internal_referral_docx,
    fill_process_note_docx,
    fill_referral_docx,
    fill_site_visit_docx,
    fill_termination_docx,
    official_c06_path,
    official_cw11_path,
    official_cw13_path,
    official_cw4a_path,
    official_cw4b_path,
)

SAMPLE = {
    '__display.organisation': 'Sebueng Itumeleng',
    '__display.personnel': 'Kelebogile Maraba',
    'household.org_household_number': '472',
    'household.date_registered': '2024-03-01',
    'caregiver.surname': 'Moholoza',
    'caregiver.name': 'Sisi Lettie',
    'caregiver.id_number': '6511010326080',
    'member.0.name': 'Thabo',
    'member.0.surname': 'Moholoza',
    'member.0.id_number': '1201015800086',
}

WORD_KEYS = (
    'monthly_report', 'educational', 'site_visit', 'exit', 'checklist',
    'content_page', 'process_note', 'termination', 'internal_referral', 'referral',
)


class RemainingWordTemplatesPresentTests(TestCase):
    def test_official_files_exist(self):
        self.assertTrue(official_c06_path().exists())
        self.assertTrue(official_cw11_path().exists())
        self.assertTrue(official_cw13_path().exists())
        self.assertTrue(official_cw4a_path().exists())
        self.assertTrue(official_cw4b_path().exists())


class RemainingWordFillTests(TestCase):
    def test_cw11_identity_row(self):
        doc = Document(BytesIO(fill_process_note_docx(SAMPLE)))
        table = doc.tables[0]
        self.assertEqual(_Cell(table.rows[1]._tr.tc_lst[0], table).text.strip(), 'Moholoza')
        self.assertEqual(_Cell(table.rows[1]._tr.tc_lst[1], table).text.strip(), 'Sisi Lettie')

    def test_cw13_identity_row(self):
        doc = Document(BytesIO(fill_termination_docx(SAMPLE)))
        table = doc.tables[0]
        self.assertEqual(_Cell(table.rows[1]._tr.tc_lst[0], table).text.strip(), 'Moholoza')

    def test_educational_family_name(self):
        text = '\n'.join(p.text for p in Document(BytesIO(fill_educational_docx(SAMPLE))).paragraphs)
        self.assertIn('Moholoza', text)

    def test_checklist_org(self):
        text = '\n'.join(p.text for p in Document(BytesIO(fill_checklist_docx(SAMPLE))).paragraphs)
        self.assertIn('Sebueng Itumeleng', text)

    def test_content_page_household_number(self):
        text = '\n'.join(p.text for p in Document(BytesIO(fill_content_page_docx(SAMPLE))).paragraphs)
        self.assertIn('472', text)

    def test_site_visit_and_exit_fill(self):
        self.assertIn('Sisi Lettie', '\n'.join(p.text for p in Document(BytesIO(fill_site_visit_docx(SAMPLE))).paragraphs))
        self.assertIn('Moholoza', '\n'.join(p.text for p in Document(BytesIO(fill_exit_docx(SAMPLE))).paragraphs))

    def test_c06_and_referrals_produce_docx(self):
        for raw in (
            fill_c06_docx(SAMPLE),
            fill_internal_referral_docx(SAMPLE),
            fill_referral_docx(SAMPLE),
        ):
            self.assertEqual(raw[:2], b'PK')


class RemainingWordPrintAndAtlasTests(TestCase):
    def setUp(self):
        Group.objects.get_or_create(name='admin')
        self.user = User.objects.create_user('remainprinter', password='x')
        self.user.groups.add(Group.objects.get(name='admin'))
        self.client = APIClient()
        self.client.credentials(HTTP_AUTHORIZATION=f'Token {Token.objects.create(user=self.user).key}')
        created = self.client.post('/api/households/', {'town': 'Westonaria'}, format='json')
        self.hid = created.data['id']
        self.token = Token.objects.get(user=self.user).key

    def test_each_remaining_form_downloads_docx(self):
        for key in WORD_KEYS:
            response = self.client.get(f'/api/print/{key}/?household_id={self.hid}&token={self.token}')
            self.assertEqual(response.status_code, 200, key)
            self.assertIn(
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                response['Content-Type'],
                key,
            )
            self.assertNotIn(b'official-payload', response.content, key)

    def test_atlas_word_geometry(self):
        self.assertIn('remaining', ATLAS_VERSION)
        for key in WORD_KEYS:
            self.assertTrue(has_geometry(key), key)
            self.assertEqual(ATLAS_FORMS[key]['geometry'], 'word', key)
            self.assertEqual(blank_info(key, 0).get('source'), 'docx', key)
            self.assertGreaterEqual(form_meta(key)['pages'], 1, key)
