"""Phase 1–2: Family Care Plan + HIV pack sheets print as filled Word."""
from io import BytesIO

from django.contrib.auth.models import Group, User
from django.test import TestCase
from docx import Document
from docx.table import _Cell
from rest_framework.authtoken.models import Token
from rest_framework.test import APIClient

from core.form_atlas import ATLAS_FORMS, fields_for, form_meta
from core.official_blanks import ATLAS_VERSION, blank_info
from core.word_forms import (
    fill_client_referral_docx,
    fill_consent_docx,
    fill_fcp_docx,
    fill_hiv_risk_docx,
    fill_hivstat_docx,
    official_fcp_path,
    official_hiv_consent_path,
    official_hiv_hts_path,
    official_hiv_pack_path,
    official_hiv_referral_path,
    official_hiv_risk_path,
)

SAMPLE = {
    '__display.organisation': 'Sebueng Itumeleng',
    'household.org_household_number': '472',
    'household.date_registered': '2024-03-01',
    'household.town': 'Westonaria',
    'caregiver.surname': 'Moholoza',
    'caregiver.name': 'Sisi Lettie',
    'caregiver.id_number': '6511010326080',
    'caregiver.cell_number': '0820000000',
    'member.0.name': 'Thabo',
    'member.0.surname': 'Moholoza',
    'member.0.id_number': '1201015800086',
    'member.0.date_of_birth': '2012-01-01',
}


class FillFamilyCarePlanTests(TestCase):
    def test_template_present(self):
        self.assertTrue(official_fcp_path().exists())

    def test_identity_header_fills(self):
        doc = Document(BytesIO(fill_fcp_docx(SAMPLE)))
        table = doc.tables[0]
        self.assertEqual(_Cell(table.rows[0]._tr.tc_lst[1], table).text.strip(), 'Moholoza')
        self.assertEqual(_Cell(table.rows[0]._tr.tc_lst[3], table).text.strip(), '472')
        self.assertEqual(_Cell(table.rows[1]._tr.tc_lst[1], table).text.strip(), '2024-03-01')


class FillHivPackSheetTests(TestCase):
    def test_split_templates_exist_beside_full_pack(self):
        self.assertTrue(official_hiv_pack_path().exists())
        self.assertTrue(official_hiv_risk_path().exists())
        self.assertTrue(official_hiv_consent_path().exists())
        self.assertTrue(official_hiv_referral_path().exists())
        self.assertTrue(official_hiv_hts_path().exists())

    def test_hiv_risk_identity_fills(self):
        doc = Document(BytesIO(fill_hiv_risk_docx(SAMPLE)))
        table = doc.tables[2]
        self.assertEqual(_Cell(table.rows[0]._tr.tc_lst[1], table).text.strip(), 'Sebueng Itumeleng')
        self.assertIn('Sisi Lettie', _Cell(table.rows[1]._tr.tc_lst[1], table).text)
        self.assertIn('Thabo', _Cell(table.rows[2]._tr.tc_lst[1], table).text)

    def test_consent_mentions_caregiver_and_child(self):
        raw = fill_consent_docx(SAMPLE)
        text = '\n'.join(p.text for p in Document(BytesIO(raw)).paragraphs)
        self.assertIn('Sisi Lettie', text)
        self.assertIn('Thabo', text)

    def test_client_referral_org_line(self):
        doc = Document(BytesIO(fill_client_referral_docx(SAMPLE)))
        org_line = next(p.text for p in doc.paragraphs if p.text.startswith('NAME OF ORGANISATION:'))
        self.assertIn('Sebueng Itumeleng', org_line)

    def test_hivstat_writes_name_boxes(self):
        doc = Document(BytesIO(fill_hivstat_docx(SAMPLE)))
        first = ''.join(_Cell(tc, doc.tables[2]).text for tc in doc.tables[2].rows[0]._tr.tc_lst[1:])
        self.assertIn('THABO', first.replace(' ', '').upper())


class PrintFcpAndHivWordTests(TestCase):
    def setUp(self):
        Group.objects.get_or_create(name='admin')
        self.user = User.objects.create_user('fcphivprinter', password='x')
        self.user.groups.add(Group.objects.get(name='admin'))
        self.client = APIClient()
        self.client.credentials(HTTP_AUTHORIZATION=f'Token {Token.objects.create(user=self.user).key}')
        created = self.client.post('/api/households/', {'town': 'Westonaria'}, format='json')
        self.hid = created.data['id']
        self.token = Token.objects.get(user=self.user).key

    def _assert_docx(self, form_key):
        response = self.client.get(f'/api/print/{form_key}/?household_id={self.hid}&token={self.token}')
        self.assertEqual(response.status_code, 200, form_key)
        self.assertIn(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            response['Content-Type'],
        )
        self.assertNotIn(b'official-payload', response.content)
        self.assertEqual(response.content[:2], b'PK')

    def test_family_care_plan_and_hiv_sheets_download_docx(self):
        for key in ('family_care_plan', 'hiv_risk', 'consent', 'client_referral', 'hivstat'):
            self._assert_docx(key)


class FcpHivAtlasTests(TestCase):
    def test_word_geometry_and_page_counts(self):
        self.assertIn('fcp', ATLAS_VERSION)
        self.assertIn('hiv', ATLAS_VERSION)
        self.assertEqual(ATLAS_FORMS['family_care_plan']['geometry'], 'word')
        self.assertEqual(ATLAS_FORMS['hiv_risk']['geometry'], 'word')
        self.assertEqual(ATLAS_FORMS['consent']['geometry'], 'word')
        self.assertEqual(ATLAS_FORMS['client_referral']['geometry'], 'word')
        self.assertEqual(ATLAS_FORMS['hivstat']['geometry'], 'word')
        self.assertEqual(form_meta('family_care_plan')['pages'], 2)
        self.assertEqual(form_meta('hiv_risk')['pages'], 3)
        self.assertEqual(form_meta('consent')['pages'], 2)
        self.assertEqual(form_meta('client_referral')['pages'], 2)
        self.assertEqual(form_meta('hivstat')['pages'], 1)
        self.assertEqual(blank_info('family_care_plan', 0).get('source'), 'docx')
        self.assertEqual(blank_info('hiv_risk', 0).get('source'), 'docx')
        self.assertEqual(fields_for('family_care_plan', 0)[0]['label'], 'Family Name')
        self.assertTrue(any(f['label'].startswith('Organisation') for f in fields_for('hiv_risk')))
