"""Phase 1–2: Official C03 prints as filled Word; blanks from Word page 1."""
from io import BytesIO

from django.contrib.auth.models import Group, User
from django.test import TestCase
from docx import Document
from docx.table import _Cell
from rest_framework.authtoken.models import Token
from rest_framework.test import APIClient

from core.form_atlas import ATLAS_FORMS, fields_for
from core.official_blanks import ATLAS_VERSION, blank_info
from core.word_forms import fill_c03_docx, official_c03_path

MPILO = {
    '__display.organisation': 'Sebueng Itumeleng',
    '__display.personnel': 'Kelebogile Maraba',
    'household.org_household_number': '472',
    'member.0.name': 'Mpilo',
    'member.0.surname': 'Khanyi',
    'member.0.id_number': '1904261049081',
}


class FillOfficialC03Tests(TestCase):
    def test_template_file_is_present(self):
        self.assertTrue(official_c03_path().exists())

    def test_child_identity_fills_the_word_header(self):
        raw = fill_c03_docx(MPILO)
        doc = Document(BytesIO(raw))
        table = doc.tables[1]
        self.assertEqual(_Cell(table.rows[0]._tr.tc_lst[1], table).text.strip(), 'Sebueng Itumeleng')
        self.assertEqual(_Cell(table.rows[1]._tr.tc_lst[1], table).text.strip(), 'Kelebogile Maraba')
        self.assertEqual(_Cell(table.rows[2]._tr.tc_lst[1], table).text.strip(), '472')
        self.assertEqual(_Cell(table.rows[3]._tr.tc_lst[1], table).text.strip(), 'Mpilo Khanyi')
        self.assertEqual(_Cell(table.rows[4]._tr.tc_lst[1], table).text.strip(), '1904261049081')


class PrintC03WordTests(TestCase):
    def setUp(self):
        Group.objects.get_or_create(name='admin')
        self.user = User.objects.create_user('c03printer', password='x')
        self.user.groups.add(Group.objects.get(name='admin'))
        self.client = APIClient()
        self.client.credentials(HTTP_AUTHORIZATION=f'Token {Token.objects.create(user=self.user).key}')

    def test_print_c03_downloads_a_docx_not_pdf_canvas(self):
        created = self.client.post('/api/households/', {'town': 'Westonaria'}, format='json')
        hid = created.data['id']
        token = Token.objects.get(user=self.user).key
        response = self.client.get(f'/api/print/c03/?household_id={hid}&token={token}')
        self.assertEqual(response.status_code, 200)
        self.assertIn(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            response['Content-Type'],
        )
        self.assertNotIn(b'official-payload', response.content)
        doc = Document(BytesIO(response.content))
        self.assertGreaterEqual(len(doc.tables), 2)


class C03WordAtlasTests(TestCase):
    def test_c03_uses_word_geometry_page1_blank(self):
        self.assertEqual(ATLAS_FORMS['c03']['geometry'], 'word')
        info = blank_info('c03', 0)
        self.assertEqual(info.get('source'), 'docx')
        self.assertIn('C03_Child_Beneficiary_Assessment.docx', info.get('source_docx') or '')
        self.assertIn('c03', ATLAS_VERSION)
        targets = {f['target'] for f in fields_for('c03')}
        self.assertIn('member.0.name', targets)
        self.assertIn('member.0.id_number', targets)
        self.assertNotIn('member.0.surname', targets)
        self.assertFalse(any(t.startswith('caregiver.') for t in targets))
