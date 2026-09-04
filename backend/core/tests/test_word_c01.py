"""Phase 1: Official C01 prints as a filled Word file, not an NPO PDF overlay."""
from io import BytesIO
from pathlib import Path

from django.contrib.auth.models import Group, User
from django.test import TestCase
from docx import Document
from rest_framework.authtoken.models import Token
from rest_framework.test import APIClient

from core.word_forms import (
    MARKED_BOX,
    fill_c01_docx,
    official_c01_path,
    _tick_choice,
)

# Ground truth from the two handwritten C01 photos the office supplied.
MOHOLOZA_VALUES = {
    'household.org_household_number': '472',
    'household.house_number': '1622',
    'household.street': 'Enkululuthweni',
    'household.town': 'Westonaria',
    'household.province': 'Gauteng',
    'household.district': 'West Rand',
    'household.municipality': 'Rand West local',
    'household.ward': '',
    '__display.personnel': 'Kelebogile Maraba',
    'caregiver.id_type': 'SA ID Number',
    'caregiver.id_number': '6511010326080',
    'caregiver.headship': 'Grand Parent Headed',
    'caregiver.name': 'Sisi Lettie',
    'caregiver.surname': 'Moholoza',
    'caregiver.nationality': 'South African',
    'caregiver.date_of_birth': '1965-11-01',
    'caregiver.sex': 'Female',
    'caregiver.race': 'African',
    'caregiver.marital_status': 'Single',
    'caregiver.disability': 'false',
    'caregiver.cell_number': '071 995 6735',
    'caregiver.home_language': 'IsiXhosa and Setswana',
    'caregiver.relationship_to_member_1': 'Grandmother',
    'member.0.id_type': 'SA ID Number',
    'member.0.id_number': '1612061437084',
    'member.0.name': 'Thandojwethu',
    'member.0.surname': 'Khanyi',
    'member.0.nationality': 'South African',
    'member.0.date_of_birth': '2016-12-06',
    'member.0.sex': 'Female',
    'member.0.race': 'African',
    'member.0.disability': 'false',
    'member.0.relationship_to_head': 'Grandchild',
    'member.1.id_type': 'SA ID Number',
    'member.1.id_number': '2407150802085',
    'member.1.name': 'Thato Buhlebendalo',
    'member.1.surname': 'Motswaledi',
    'member.1.nationality': 'South African',
    'member.1.date_of_birth': '2024-07-15',
    'member.1.sex': 'Female',
    'member.1.race': 'African',
    'member.1.disability': 'false',
    'member.1.relationship_to_head': 'Grandchild',
    'member.2.id_type': 'SA ID Number',
    'member.2.id_number': '2212280261081',
    'member.2.name': 'Onkargbetswe Busisiwe',
    'member.2.surname': 'Kgoasi',
    'member.2.nationality': 'South African',
    'member.2.date_of_birth': '2022-12-28',
    'member.2.sex': 'Female',
    'member.2.race': 'African',
    'member.2.disability': 'false',
    'member.2.relationship_to_head': 'Grandchild',
    'member.3.id_type': 'SA ID Number',
    'member.3.id_number': '0105281329083',
    'member.3.name': 'Botshelo Masego',
    'member.3.surname': 'Motswaledi',
    'member.3.nationality': 'South African',
    'member.3.date_of_birth': '2001-05-28',
    'member.3.sex': 'Female',
    'member.3.race': 'African',
    'member.3.disability': 'false',
    'member.3.relationship_to_head': 'Grandchild',
}


class TickChoiceTests(TestCase):
    def test_grand_parent_does_not_mark_parent(self):
        text = (
            'Parent Headed  ☐     Grand Parent Headed  ☐     Youth Headed  ☐\n'
            'Child Headed  ☐     Relative Headed  ☐     Other  ☐'
        )
        marked = _tick_choice(
            text, 'Grand Parent Headed',
            options=(
                'Grand Parent Headed', 'Parent Headed', 'Youth Headed',
                'Child Headed', 'Relative Headed', 'Other',
            ),
        )
        self.assertIn(f'Grand Parent Headed  {MARKED_BOX}', marked)
        self.assertIn('Parent Headed  ☐', marked)


class FillOfficialC01Tests(TestCase):
    def test_template_file_is_present(self):
        self.assertTrue(official_c01_path().exists())

    def test_moholoza_photo_values_fill_the_word_template(self):
        raw = fill_c01_docx(MOHOLOZA_VALUES)
        doc = Document(BytesIO(raw))
        page1 = '\n'.join(c.text for row in doc.tables[1].rows for c in row.cells)
        self.assertIn('472', page1)
        self.assertIn('Westonaria', page1)
        self.assertIn('Enkululuthweni', page1)
        self.assertIn('Sisi Lettie', page1)
        self.assertIn('Moholoza', page1)
        self.assertIn('6511010326080', page1)
        self.assertIn('Thandojwethu', page1)
        self.assertIn('Khanyi', page1)
        self.assertIn(f'Grand Parent Headed  {MARKED_BOX}', page1)
        self.assertIn(f'Female  {MARKED_BOX}', page1)
        member2 = '\n'.join(c.text for row in doc.tables[3].rows for c in row.cells)
        self.assertIn('Thato Buhlebendalo', member2)
        self.assertIn('Motswaledi', member2)
        self.assertIn('2407150802085', member2)
        member3 = '\n'.join(c.text for row in doc.tables[4].rows for c in row.cells)
        self.assertIn('Onkargbetswe Busisiwe', member3)
        self.assertIn('2212280261081', member3)
        member4 = '\n'.join(c.text for row in doc.tables[5].rows for c in row.cells)
        self.assertIn('Botshelo Masego', member4)
        self.assertIn('0105281329083', member4)


class PrintC01WordTests(TestCase):
    def setUp(self):
        Group.objects.get_or_create(name='admin')
        self.user = User.objects.create_user(
            'OrphanCoordinator', password='x', is_staff=True, is_superuser=True,
        )
        self.user.groups.add(Group.objects.get(name='admin'))
        self.token = Token.objects.create(user=self.user)
        self.client = APIClient()
        self.client.credentials(HTTP_AUTHORIZATION=f'Token {self.token.key}')

    def test_print_c01_downloads_a_docx_not_pdf_canvas(self):
        created = self.client.post('/api/households/', {
            'org_household_number': '472',
            'town': 'Westonaria',
            'province': 'Gauteng',
            'district': 'West Rand',
            'municipality': 'Rand West local',
            'street': 'Enkululuthweni',
            'house_number': '1622',
        }, format='json')
        self.assertEqual(created.status_code, 201, created.data)
        hid = created.data['id']
        cg = self.client.post('/api/caregivers/', {
            'household': hid,
            'name': 'Sisi Lettie',
            'surname': 'Moholoza',
            'id_number': '6511010326080',
            'sex': 'Female',
            'date_of_birth': '1965-11-01',
            'surname_confirmed': True,
            'id_number_confirmed': True,
            'date_of_birth_confirmed': True,
        }, format='json')
        self.assertEqual(cg.status_code, 201, cg.data)
        token = self.token.key
        resp = self.client.get(f'/api/print/c01/?household_id={hid}&token={token}')
        self.assertEqual(resp.status_code, 200)
        self.assertIn(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resp['Content-Type'],
        )
        self.assertNotIn(b'official-payload', resp.content)
        doc = Document(BytesIO(resp.content))
        page1 = '\n'.join(c.text for row in doc.tables[1].rows for c in row.cells)
        self.assertIn('Westonaria', page1)
        self.assertIn('Moholoza', page1)
        self.assertIn('6511010326080', page1)
