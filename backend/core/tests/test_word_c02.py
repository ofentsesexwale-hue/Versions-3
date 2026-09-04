"""Phase 1–2: Official C02 prints as filled Word; blanks from the Word template."""
from io import BytesIO
from pathlib import Path

from django.contrib.auth.models import Group, User
from django.test import TestCase
from docx import Document
from docx.table import _Cell
from rest_framework.authtoken.models import Token
from rest_framework.test import APIClient

from core.form_atlas import ATLAS_FORMS, fields_for
from core.official_blanks import ATLAS_VERSION, blank_info
from core.word_forms import fill_c02_docx, official_c02_path

MOHOLOZA_C02 = {
    '__display.organisation': 'Sebueng Itumeleng',
    '__display.personnel': 'Kelebogile Maraba',
    'household.org_household_number': '472',
    'caregiver.name': 'Sisi Lettie',
    'caregiver.surname': 'Moholoza',
    'caregiver.id_number': '6511010326080',
}


class FillOfficialC02Tests(TestCase):
    def test_template_file_is_present(self):
        self.assertTrue(official_c02_path().exists())

    def test_moholoza_identity_fills_the_word_header(self):
        raw = fill_c02_docx(MOHOLOZA_C02)
        doc = Document(BytesIO(raw))
        table = doc.tables[1]
        self.assertEqual(_Cell(table.rows[0]._tr.tc_lst[1], table).text.strip(), 'Sebueng Itumeleng')
        self.assertEqual(_Cell(table.rows[1]._tr.tc_lst[1], table).text.strip(), 'Kelebogile Maraba')
        self.assertEqual(_Cell(table.rows[2]._tr.tc_lst[1], table).text.strip(), '472')
        self.assertEqual(_Cell(table.rows[3]._tr.tc_lst[1], table).text.strip(), 'Sisi Lettie Moholoza')
        self.assertEqual(_Cell(table.rows[4]._tr.tc_lst[1], table).text.strip(), '6511010326080')


class PrintC02WordTests(TestCase):
    def setUp(self):
        Group.objects.get_or_create(name='admin')
        self.user = User.objects.create_user('c02printer', password='x')
        self.user.groups.add(Group.objects.get(name='admin'))
        self.client = APIClient()
        self.client.credentials(HTTP_AUTHORIZATION=f'Token {Token.objects.create(user=self.user).key}')

    def test_print_c02_downloads_a_docx_not_pdf_canvas(self):
        created = self.client.post('/api/households/', {'town': 'Westonaria'}, format='json')
        hid = created.data['id']
        self.client.patch(
            f'/api/households/{hid}/',
            {'org_household_number': '472'},
            format='json',
        )
        token = Token.objects.get(user=self.user).key
        response = self.client.get(f'/api/print/c02/?household_id={hid}&token={token}')
        self.assertEqual(response.status_code, 200)
        self.assertIn(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            response['Content-Type'],
        )
        self.assertNotIn(b'official-payload', response.content)
        doc = Document(BytesIO(response.content))
        self.assertGreaterEqual(len(doc.tables), 2)


class C02WordAtlasTests(TestCase):
    def test_c02_uses_word_geometry_and_blank(self):
        self.assertEqual(ATLAS_FORMS['c02']['geometry'], 'word')
        info = blank_info('c02', 0)
        self.assertEqual(info.get('source'), 'docx')
        self.assertIn('C02_Adult_Assessment_Form.docx', info.get('source_docx') or '')
        self.assertIn('word-c01-c02', ATLAS_VERSION)
        targets = {f['target'] for f in fields_for('c02')}
        self.assertIn('caregiver.name', targets)
        self.assertIn('caregiver.id_number', targets)
        self.assertIn('household.org_household_number', targets)
        self.assertNotIn('caregiver.nationality', targets)
