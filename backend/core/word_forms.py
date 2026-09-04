"""Fill official DSD Word templates. Do not overlay the NPO PDF.

The NPO case-management PDF is a file-order guide only. Print for C01–C03 and
CW 05 writes into the Official Word templates under ``docs/official/dsd-source/``.
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.table import _Cell

REPO_ROOT = Path(__file__).resolve().parents[2]
OFFICIAL_C01 = REPO_ROOT / 'docs' / 'official' / 'dsd-source' / 'Official_C01_Template.docx'
OFFICIAL_C02 = REPO_ROOT / 'docs' / 'official' / 'dsd-source' / 'C02_Adult_Assessment_Form.docx'
OFFICIAL_C03 = REPO_ROOT / 'docs' / 'official' / 'dsd-source' / 'C03_Child_Beneficiary_Assessment.docx'
OFFICIAL_CW05 = REPO_ROOT / 'docs' / 'official' / 'dsd-source' / 'CW_05_Intake_Form_28082019.docx'

EMPTY_BOX = '☐'
MARKED_BOX = '☑'


def official_c01_path() -> Path:
    return OFFICIAL_C01


def official_c02_path() -> Path:
    return OFFICIAL_C02


def official_c03_path() -> Path:
    return OFFICIAL_C03


def official_cw05_path() -> Path:
    return OFFICIAL_CW05


# Alias used by print/scan code (atlas key is ``intake``).
official_intake_path = official_cw05_path


def _tc_cell(table, row_index: int, cell_index: int) -> _Cell:
    """Return the real cell at (row, col) via XML — row.cells lies on C02 merges."""
    return _Cell(table.rows[row_index]._tr.tc_lst[cell_index], table)


def _clear_cell(cell, text: str = ''):
    """Replace every paragraph in a cell with a single run of ``text``."""
    text = '' if text is None else str(text)
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return
    first = paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    if first.runs:
        first.runs[0].text = text
    else:
        first.add_run(text)
    for para in paragraphs[1:]:
        para._element.getparent().remove(para._element)


def _set_value_row(row, value: str, value_col: int = 1):
    """Write into the first value cell; mirror across merged duplicates."""
    cells = row.cells
    if value_col >= len(cells):
        return
    _clear_cell(cells[value_col], value)
    # Merged rows often repeat the same cell object; unique cells only.
    seen = {id(cells[value_col]._tc)}
    for cell in cells[value_col + 1:]:
        if id(cell._tc) in seen:
            continue
        # Only overwrite blank lookalikes (merged copies), not a second label.
        if (cell.text or '').strip() in ('', (cells[value_col].text or '').strip()):
            _clear_cell(cell, value)
            seen.add(id(cell._tc))


def _tick_choice(cell_or_text, chosen: str, options: tuple[str, ...] | None = None) -> str:
    """Mark the matching option with ☑ and leave the others ☐."""
    text = cell_or_text.text if hasattr(cell_or_text, 'text') else (cell_or_text or '')
    chosen = (chosen or '').strip()
    if not chosen:
        return text
    working = text.replace(MARKED_BOX, EMPTY_BOX)
    # Longest option first so "Grand Parent Headed" wins over "Parent Headed".
    option_names = sorted(options or (chosen,), key=len, reverse=True)
    match = next((opt for opt in option_names if opt.lower() == chosen.lower()), chosen)
    # Replace only the first "Match  ☐" / "Match ☐" occurrence for that option.
    for gap in ('  ', ' '):
        needle = f'{match}{gap}{EMPTY_BOX}'
        if needle in working:
            return working.replace(needle, f'{match}{gap}{MARKED_BOX}', 1)
    return working


HEADSHIP_OPTIONS = (
    'Grand Parent Headed', 'Parent Headed', 'Youth Headed',
    'Child Headed', 'Relative Headed', 'Other',
)
ID_TYPE_OPTIONS = ('SA ID Number', 'Passport Number', 'Permit')
SEX_OPTIONS = ('Male', 'Female')
RACE_OPTIONS = ('African', 'White', 'Coloured', 'Indian')
NATIONALITY_OPTIONS = ('South African', 'Other')
MARITAL_OPTIONS = ('Married', 'Divorced', 'Widowed', 'Single', 'Cohabiting', 'Separate')
DISABILITY_OPTIONS = ('No', 'Yes')


def _apply_ticks(row, chosen: str, start_col: int = 1, options: tuple[str, ...] | None = None):
    cells = row.cells
    if start_col >= len(cells):
        return
    marked = _tick_choice(cells[start_col], chosen, options=options)
    seen = set()
    for cell in cells[start_col:]:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        if EMPTY_BOX in (cell.text or '') or MARKED_BOX in (cell.text or ''):
            _clear_cell(cell, marked)


def _member_payload(values: dict, slot: int) -> dict:
    p = f'member.{slot}.'
    return {
        'id_type': values.get(p + 'id_type') or '',
        'id_number': values.get(p + 'id_number') or '',
        'name': values.get(p + 'name') or '',
        'surname': values.get(p + 'surname') or '',
        'known_as': values.get(p + 'known_as') or '',
        'nationality': values.get(p + 'nationality') or '',
        'date_of_birth': values.get(p + 'date_of_birth') or '',
        'sex': values.get(p + 'sex') or '',
        'race': values.get(p + 'race') or '',
        'disability': values.get(p + 'disability') or '',
        'disability_description': values.get(p + 'disability_description') or '',
        'date_joined': values.get(p + 'date_joined') or '',
        'relationship_to_head': values.get(p + 'relationship_to_head') or '',
    }


def _fill_member_table(table, person: dict):
    """Tables 3/4/5: Add member 2/3/4 — label | value."""
    if not any(person.values()):
        return
    _apply_ticks(table.rows[1], person.get('id_type') or 'SA ID Number', options=ID_TYPE_OPTIONS)
    _set_value_row(table.rows[2], person.get('id_number') or '', 1)
    _set_value_row(table.rows[3], person.get('name') or '', 1)
    _set_value_row(table.rows[4], person.get('surname') or '', 1)
    _set_value_row(table.rows[5], person.get('known_as') or '', 1)
    _set_value_row(table.rows[6], person.get('nationality') or '', 1)
    _set_value_row(table.rows[7], person.get('date_of_birth') or '', 1)
    _apply_ticks(table.rows[8], person.get('sex') or '', options=SEX_OPTIONS)
    _apply_ticks(table.rows[9], person.get('race') or '', options=RACE_OPTIONS)
    disability = person.get('disability') or ''
    if disability in ('true', 'True', True, 'yes', 'Yes'):
        label = 'Yes'
    elif disability in ('false', 'False', False, 'no', 'No'):
        label = 'No'
    else:
        label = disability
    marked = _tick_choice(table.rows[10].cells[1], label, options=DISABILITY_OPTIONS)
    if person.get('disability_description'):
        marked = marked.replace('____________________', person['disability_description'])
        marked = marked.replace('___________________', person['disability_description'])
    _clear_cell(table.rows[10].cells[1], marked)
    _set_value_row(table.rows[11], person.get('date_joined') or '', 1)
    _set_value_row(table.rows[12], person.get('relationship_to_head') or '', 1)


def _fill_page1(table, values: dict):
    """Main household + caregiver + member 1 table."""
    _set_value_row(table.rows[1], values.get('household.org_household_number') or '', 1)
    _set_value_row(table.rows[1], values.get('household.province') or '', 3)
    _set_value_row(table.rows[2], values.get('household.house_number') or '', 1)
    _set_value_row(table.rows[2], values.get('household.district') or '', 3)
    _set_value_row(table.rows[3], values.get('household.street') or '', 1)
    _set_value_row(table.rows[3], values.get('household.municipality') or '', 3)
    _set_value_row(table.rows[4], values.get('household.town') or '', 1)
    _set_value_row(table.rows[4], values.get('household.ward') or '', 3)

    personnel = values.get('__display.personnel') or values.get('household.personnel') or ''
    # Personnel row: label in col0, value spans the rest.
    row = table.rows[5]
    if len(row.cells) > 1:
        _clear_cell(row.cells[1], personnel)

    _apply_ticks(table.rows[7], values.get('caregiver.id_type') or 'SA ID Number', options=ID_TYPE_OPTIONS)
    _set_value_row(table.rows[8], values.get('caregiver.id_number') or '', 1)
    _set_value_row(table.rows[9], values.get('caregiver.organisation_beneficiary_number') or '', 1)
    _apply_ticks(table.rows[10], values.get('caregiver.headship') or '', options=HEADSHIP_OPTIONS)
    _set_value_row(table.rows[11], values.get('caregiver.name') or '', 1)
    _set_value_row(table.rows[12], values.get('caregiver.surname') or '', 1)
    _set_value_row(table.rows[13], values.get('caregiver.known_as') or '', 1)
    _apply_ticks(table.rows[14], values.get('caregiver.nationality') or '', options=NATIONALITY_OPTIONS)
    _set_value_row(table.rows[15], values.get('caregiver.date_of_birth') or '', 1)
    _apply_ticks(table.rows[16], values.get('caregiver.sex') or '', options=SEX_OPTIONS)
    _apply_ticks(table.rows[17], values.get('caregiver.race') or '', options=RACE_OPTIONS)
    _apply_ticks(table.rows[18], values.get('caregiver.marital_status') or '', options=MARITAL_OPTIONS)

    disability = values.get('caregiver.disability') or ''
    if disability in ('true', True, 'yes', 'Yes'):
        dlabel = 'Yes'
    elif disability in ('false', False, 'no', 'No'):
        dlabel = 'No'
    else:
        dlabel = disability
    marked = _tick_choice(table.rows[19].cells[1], dlabel, options=DISABILITY_OPTIONS)
    describe = values.get('caregiver.disability_description') or ''
    if describe:
        marked = marked.replace('____________________', describe)
    _clear_cell(table.rows[19].cells[1], marked)

    _set_value_row(table.rows[20], values.get('caregiver.cell_number') or '', 1)
    _set_value_row(table.rows[21], values.get('caregiver.home_language') or '', 1)
    _set_value_row(table.rows[22], values.get('caregiver.date_joined') or '', 1)

    # Relationship to members line.
    rels = []
    for i in range(4):
        rel = values.get(f'member.{i}.relationship_to_head') or ''
        if i == 0 and values.get('caregiver.relationship_to_member_1'):
            rel = values.get('caregiver.relationship_to_member_1') or rel
        rels.append(rel)
    rel_line = (
        f"Member 1 {rels[0] or '_____________'}  "
        f"Member 2 {rels[1] or '______________'}  "
        f"Member 3 {rels[2] or '_________________'}   "
        f"Member 4 {rels[3] or '______________'}"
    )
    _clear_cell(table.rows[23].cells[1], rel_line)

    # Member 1 block (same table).
    m0 = _member_payload(values, 0)
    _apply_ticks(table.rows[25], m0.get('id_type') or ('SA ID Number' if m0.get('id_number') else ''), options=ID_TYPE_OPTIONS)
    _set_value_row(table.rows[26], m0.get('id_number') or '', 1)
    _set_value_row(table.rows[27], m0.get('name') or '', 1)
    _set_value_row(table.rows[28], m0.get('surname') or '', 1)
    _set_value_row(table.rows[29], m0.get('known_as') or '', 1)
    _set_value_row(table.rows[30], m0.get('nationality') or '', 1)
    _set_value_row(table.rows[31], m0.get('date_of_birth') or '', 1)
    _apply_ticks(table.rows[32], m0.get('sex') or '', options=SEX_OPTIONS)
    _apply_ticks(table.rows[33], m0.get('race') or '', options=RACE_OPTIONS)
    disability = m0.get('disability') or ''
    if disability in ('true', True, 'yes', 'Yes'):
        dlabel = 'Yes'
    elif disability in ('false', False, 'no', 'No'):
        dlabel = 'No'
    else:
        dlabel = disability
    marked = _tick_choice(table.rows[34].cells[1], dlabel, options=DISABILITY_OPTIONS)
    if m0.get('disability_description'):
        marked = marked.replace('____________________', m0['disability_description'])
    _clear_cell(table.rows[34].cells[1], marked)
    _set_value_row(table.rows[35], m0.get('date_joined') or '', 1)
    _set_value_row(table.rows[36], m0.get('relationship_to_head') or '', 1)


def fill_c01_docx(values: dict | None = None, template_path: Path | None = None) -> bytes:
    """Return a filled Official C01 .docx (bytes) from atlas-style values."""
    path = Path(template_path or official_c01_path())
    if not path.exists():
        raise FileNotFoundError(f'Official C01 template missing: {path}')
    values = values or {}
    doc = Document(str(path))
    if len(doc.tables) < 6:
        raise ValueError('Official C01 template does not have the expected tables')
    _fill_page1(doc.tables[1], values)
    _fill_member_table(doc.tables[3], _member_payload(values, 1))
    _fill_member_table(doc.tables[4], _member_payload(values, 2))
    _fill_member_table(doc.tables[5], _member_payload(values, 3))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _adult_display_name(values: dict) -> str:
    name = (values.get('caregiver.name') or '').strip()
    surname = (values.get('caregiver.surname') or '').strip()
    return ' '.join(part for part in (name, surname) if part)


def _fill_c02_header(table, values: dict):
    """Identity header on C02_Adult_Assessment_Form.docx (table 1)."""
    org = (
        values.get('__display.organisation')
        or values.get('organisation.name')
        or ''
    )
    personnel = values.get('__display.personnel') or ''
    org_hh = values.get('household.org_household_number') or ''
    full_name = _adult_display_name(values)
    id_number = values.get('caregiver.id_number') or ''

    _clear_cell(_tc_cell(table, 0, 1), org)
    _clear_cell(_tc_cell(table, 1, 1), personnel)
    _clear_cell(_tc_cell(table, 2, 1), org_hh)
    _clear_cell(_tc_cell(table, 3, 1), full_name)
    _clear_cell(_tc_cell(table, 4, 1), id_number)


def fill_c02_docx(values: dict | None = None, template_path: Path | None = None) -> bytes:
    """Return a filled Official C02 .docx (bytes) from atlas-style values.

    Assessment Yes/No/DK grids stay blank for staff to tick on paper / later
    atlas work. Identity header is filled from the caregiver file.
    """
    path = Path(template_path or official_c02_path())
    if not path.exists():
        raise FileNotFoundError(f'Official C02 template missing: {path}')
    values = values or {}
    doc = Document(str(path))
    if len(doc.tables) < 2:
        raise ValueError('Official C02 template does not have the expected header table')
    _fill_c02_header(doc.tables[1], values)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _child_display_name(values: dict, slot: int = 0) -> str:
    p = f'member.{slot}.'
    name = (values.get(p + 'name') or '').strip()
    surname = (values.get(p + 'surname') or '').strip()
    return ' '.join(part for part in (name, surname) if part)


def _resolve_member_slot(values: dict, requested=None) -> int:
    if requested is not None and str(requested).strip() != '':
        try:
            return max(0, int(requested))
        except (TypeError, ValueError):
            pass
    for i in range(8):
        p = f'member.{i}.'
        if values.get(p + 'id_number') or values.get(p + 'name') or values.get(p + 'surname'):
            return i
    return 0


def _fill_c03_header(table, values: dict, slot: int = 0):
    """Identity header on C03_Child_Beneficiary_Assessment.docx (table 1, page 1)."""
    p = f'member.{slot}.'
    org = (
        values.get('__display.organisation')
        or values.get('organisation.name')
        or ''
    )
    personnel = (
        values.get('__display.personnel')
        or values.get('__display.cycw_name')
        or ''
    )
    org_hh = values.get('household.org_household_number') or ''
    full_name = _child_display_name(values, slot)
    id_number = values.get(p + 'id_number') or ''

    _clear_cell(_tc_cell(table, 0, 1), org)
    _clear_cell(_tc_cell(table, 1, 1), personnel)
    _clear_cell(_tc_cell(table, 2, 1), org_hh)
    _clear_cell(_tc_cell(table, 3, 1), full_name)
    _clear_cell(_tc_cell(table, 4, 1), id_number)


def fill_c03_docx(
    values: dict | None = None,
    template_path: Path | None = None,
    member_slot=None,
) -> bytes:
    """Return a filled Official C03 .docx (bytes) for one child beneficiary.

    Only page 1 of the Word template is used (LibreOffice exports one page;
    any trailing blank page in Word is ignored). Assessment ticks stay blank.
    """
    path = Path(template_path or official_c03_path())
    if not path.exists():
        raise FileNotFoundError(f'Official C03 template missing: {path}')
    values = values or {}
    slot = _resolve_member_slot(values, member_slot)
    doc = Document(str(path))
    if len(doc.tables) < 2:
        raise ValueError('Official C03 template does not have the expected header table')
    _fill_c03_header(doc.tables[1], values, slot=slot)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fill_cw05_identity(table, values: dict):
    """Identity block on CW_05_Intake_Form_28082019.docx (table 0)."""
    _clear_cell(_tc_cell(table, 1, 1), values.get('household.org_household_number') or '')
    _clear_cell(_tc_cell(table, 3, 0), values.get('caregiver.surname') or '')
    _clear_cell(_tc_cell(table, 3, 1), values.get('caregiver.name') or '')
    id_or_dob = values.get('caregiver.id_number') or values.get('caregiver.date_of_birth') or ''
    _clear_cell(_tc_cell(table, 3, 2), id_or_dob)


def fill_cw05_docx(values: dict | None = None, template_path: Path | None = None) -> bytes:
    """Return a filled Official CW 05 Intake .docx (bytes).

    Fills intake ref + primary client identity. Caregiver-of-child rows,
    narratives, risk ticks, and consent stay blank for paper / later atlas.
    """
    path = Path(template_path or official_cw05_path())
    if not path.exists():
        raise FileNotFoundError(f'Official CW 05 template missing: {path}')
    values = values or {}
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError('Official CW 05 template does not have the expected tables')
    _fill_cw05_identity(doc.tables[0], values)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Atlas / print key for CW 05 is ``intake``.
fill_intake_docx = fill_cw05_docx


def values_from_household(household) -> dict:
    """Atlas-style values plus free-text fields the Word sheet needs."""
    from .form_io import values_for_household
    from .models import Organisation

    values = values_for_household(household)
    org = Organisation.objects.first()
    if org and not values.get('__display.organisation'):
        values['__display.organisation'] = org.name or 'Sebueng Itumeleng'
    elif not values.get('__display.organisation'):
        values['__display.organisation'] = 'Sebueng Itumeleng'
    cg = getattr(household, 'caregiver', None)
    if cg:
        mapping = {
            'caregiver.marital_status': 'marital_status',
            'caregiver.cell_number': 'cell_number',
            'caregiver.home_language': 'home_language',
            'caregiver.organisation_beneficiary_number': 'organisation_beneficiary_number',
            'caregiver.nationality': 'nationality',
            'caregiver.headship': 'headship_type',
        }
        for key, attr in mapping.items():
            if values.get(key):
                continue
            raw = getattr(cg, attr, '') if hasattr(cg, attr) else ''
            if raw is True:
                values[key] = 'true'
            elif raw is False:
                values[key] = 'false'
            elif raw is not None and raw != '':
                values[key] = str(raw)
    if household and not values.get('household.municipality'):
        values['household.municipality'] = getattr(household, 'municipality', '') or ''
    return values
